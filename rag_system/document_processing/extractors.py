"""
文档文本提取器
"""
import logging
import mimetypes
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Union
import re

from ..utils.exceptions import DocumentError

logger = logging.getLogger(__name__)


class BaseTextExtractor(ABC):
    """文本提取器基类"""
    
    @abstractmethod
    def extract(self, file_path: Union[str, Path]) -> str:
        """提取文本内容"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        pass
    
    def validate_file(self, file_path: Union[str, Path]) -> bool:
        """验证文件格式"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise DocumentError(f"不是有效文件: {file_path}")
        
        extension = file_path.suffix.lower()
        if extension not in self.get_supported_extensions():
            raise DocumentError(f"不支持的文件格式: {extension}")
        
        return True


class TxtExtractor(BaseTextExtractor):
    """TXT文件提取器"""
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """提取TXT文件内容"""
        self.validate_file(file_path)
        
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.debug(f"成功使用 {encoding} 编码读取文件: {file_path}")
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentError(f"无法解码文件: {file_path}")
            
        except Exception as e:
            logger.error(f"提取TXT文件失败: {file_path}, 错误: {str(e)}")
            raise DocumentError(f"提取TXT文件失败: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['.txt']


class MarkdownExtractor(BaseTextExtractor):
    """Markdown文件提取器"""
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """提取Markdown文件内容"""
        self.validate_file(file_path)
        
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    # 简单的Markdown清理
                    cleaned_content = self._clean_markdown(content)
                    logger.debug(f"成功提取Markdown文件: {file_path}")
                    return cleaned_content.strip()
                    
                except UnicodeDecodeError:
                    continue
            
            raise DocumentError(f"无法解码Markdown文件: {file_path}")
            
        except Exception as e:
            logger.error(f"提取Markdown文件失败: {file_path}, 错误: {str(e)}")
            raise DocumentError(f"提取Markdown文件失败: {str(e)}")
    
    def _clean_markdown(self, content: str) -> str:
        """清理Markdown格式"""
        # 移除标题标记
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        
        # 移除粗体和斜体标记
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        content = re.sub(r'__(.*?)__', r'\1', content)
        content = re.sub(r'_(.*?)_', r'\1', content)
        
        # 移除链接，保留文本
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        
        # 移除图片
        content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', content)
        
        # 移除代码块标记
        content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        # 移除引用标记
        content = re.sub(r'^>\s+', '', content, flags=re.MULTILINE)
        
        # 移除列表标记
        content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)
        
        # 清理多余的空行
        content = re.sub(r'\n\s*\n', '\n\n', content)
        
        return content
    
    def get_supported_extensions(self) -> List[str]:
        return ['.md', '.markdown']


class PDFExtractor(BaseTextExtractor):
    """PDF文件提取器"""
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖库"""
        try:
            import PyPDF2
            self.PyPDF2 = PyPDF2
        except ImportError:
            try:
                import pypdf
                self.PyPDF2 = pypdf
            except ImportError:
                raise DocumentError("PDF处理需要安装PyPDF2或pypdf: pip install PyPDF2")
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """提取PDF文件内容"""
        self.validate_file(file_path)
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = self.PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise DocumentError("PDF文件没有页面")
                
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                        logger.debug(f"提取PDF第{page_num + 1}页内容")
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num + 1}页失败: {str(e)}")
                        continue
                
                if not text_content:
                    raise DocumentError("PDF文件中没有可提取的文本内容")
                
                full_text = '\n\n'.join(text_content)
                logger.info(f"成功提取PDF文件: {file_path}, 页数: {len(pdf_reader.pages)}")
                return full_text.strip()
                
        except Exception as e:
            logger.error(f"提取PDF文件失败: {file_path}, 错误: {str(e)}")
            raise DocumentError(f"提取PDF文件失败: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['.pdf']


class DocxExtractor(BaseTextExtractor):
    """DOCX文件提取器"""
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖库"""
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise DocumentError("DOCX处理需要安装python-docx: pip install python-docx")
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """提取DOCX文件内容"""
        self.validate_file(file_path)
        
        try:
            doc = self.docx.Document(file_path)
            
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise DocumentError("DOCX文件中没有可提取的文本内容")
            
            full_text = '\n\n'.join(text_content)
            logger.info(f"成功提取DOCX文件: {file_path}")
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"提取DOCX文件失败: {file_path}, 错误: {str(e)}")
            raise DocumentError(f"提取DOCX文件失败: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx']


class TextExtractorFactory:
    """文本提取器工厂"""
    
    def __init__(self):
        self._extractors: Dict[str, BaseTextExtractor] = {}
        self._register_default_extractors()
    
    def _register_default_extractors(self):
        """注册默认提取器"""
        try:
            # TXT提取器
            txt_extractor = TxtExtractor()
            for ext in txt_extractor.get_supported_extensions():
                self._extractors[ext] = txt_extractor
            
            # Markdown提取器
            md_extractor = MarkdownExtractor()
            for ext in md_extractor.get_supported_extensions():
                self._extractors[ext] = md_extractor
            
            # PDF提取器
            try:
                pdf_extractor = PDFExtractor()
                for ext in pdf_extractor.get_supported_extensions():
                    self._extractors[ext] = pdf_extractor
            except DocumentError as e:
                logger.warning(f"PDF提取器注册失败: {str(e)}")
            
            # DOCX提取器
            try:
                docx_extractor = DocxExtractor()
                for ext in docx_extractor.get_supported_extensions():
                    self._extractors[ext] = docx_extractor
            except DocumentError as e:
                logger.warning(f"DOCX提取器注册失败: {str(e)}")
                
        except Exception as e:
            logger.error(f"注册默认提取器失败: {str(e)}")
    
    def register_extractor(self, extension: str, extractor: BaseTextExtractor):
        """注册自定义提取器"""
        if not extension.startswith('.'):
            extension = '.' + extension
        
        extension = extension.lower()
        self._extractors[extension] = extractor
        logger.info(f"注册文本提取器: {extension}")
    
    def get_extractor(self, file_path: Union[str, Path]) -> Optional[BaseTextExtractor]:
        """获取文件对应的提取器"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        return self._extractors.get(extension)
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """提取文件文本内容"""
        extractor = self.get_extractor(file_path)
        
        if not extractor:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            supported_formats = list(self._extractors.keys())
            raise DocumentError(
                f"不支持的文件格式: {extension}. "
                f"支持的格式: {', '.join(supported_formats)}"
            )
        
        return extractor.extract(file_path)
    
    def get_supported_formats(self) -> List[str]:
        """获取所有支持的文件格式"""
        return list(self._extractors.keys())
    
    def detect_file_type(self, file_path: Union[str, Path]) -> Optional[str]:
        """检测文件类型"""
        file_path = Path(file_path)
        
        # 首先通过扩展名检测
        extension = file_path.suffix.lower()
        if extension in self._extractors:
            return extension
        
        # 通过MIME类型检测
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            mime_to_ext = {
                'text/plain': '.txt',
                'text/markdown': '.md',
                'application/pdf': '.pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
            }
            
            detected_ext = mime_to_ext.get(mime_type)
            if detected_ext and detected_ext in self._extractors:
                return detected_ext
        
        return None
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """检查文件是否支持"""
        return self.detect_file_type(file_path) is not None